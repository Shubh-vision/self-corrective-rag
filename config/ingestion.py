import streamlit as st
from langchain_community.document_loaders import (
    PyPDFLoader,
    WebBaseLoader,
    CSVLoader,
    UnstructuredWordDocumentLoader
)
# For PDF in-memory:
import pypdf 
# For DOCX in-memory:
from docx import Document as DocxReader
import pandas as pd
from langchain_core.documents import Document

# ================= LOAD FILE / URL / TEXT =================
def load_source(uploaded_file=None, url=None, raw_text=None):
    docs = []


    # ================= PRIORITY RULE =================
    if uploaded_file:
        name = uploaded_file.name.lower()
        
        # --- 1. PDF Handling (In-Memory) ---
        if name.endswith(".pdf"):
            pdf_reader = pypdf.PdfReader(uploaded_file)
            pdf_text = ""
            for page in pdf_reader.pages:
                pdf_text += page.extract_text() + "\n"
            docs.append(Document(page_content=pdf_text, metadata={"source": name}))



        # --- 2. DOCX Handling (In-Memory) ---
        elif name.endswith(".docx"):
            doc = DocxReader(uploaded_file)
            docx_text = "\n".join([para.text for para in doc.paragraphs])
            docs.append(Document(page_content=docx_text, metadata={"source": name}))

            

        # --- 3. CSV Handling (Temp-file bypass) ---
        elif name.endswith(".csv"):
            # CSVLoader requires a path, but we can process the string directly 
            # or use a StringIO workaround if you prefer the LangChain Loader
            df = pd.read_csv(uploaded_file)
            content = df.to_string()
            docs.append(Document(page_content=content, metadata={"source": name}))



     # 2. URL (only if file not given)
    elif url:
        docs.extend(WebBaseLoader(url).load())

    # 3. RAW TEXT (only if others not given)
    elif raw_text:
        docs.append(Document(page_content=raw_text))

    return docs              


def get_docs(uploaded_file=None, url=None, raw_text=None):
    return load_source(uploaded_file, url, raw_text)
